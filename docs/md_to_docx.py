import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def convert_md_to_docx(md_file, docx_file):
    # 创建Word文档
    doc = docx.Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # 转换处理
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 处理标题
        if line.startswith('# '):
            add_heading(doc, line[2:], 0)
        elif line.startswith('## '):
            add_heading(doc, line[3:], 1)
        elif line.startswith('### '):
            add_heading(doc, line[4:], 2)
        # 处理列表
        elif line.startswith('- '):
            add_list_item(doc, line[2:])
        # 处理图片
        elif '![' in line:
            add_image_placeholder(doc, line)
        # 普通段落
        else:
            add_paragraph(doc, line)
    
    # 保存Word文档
    doc.save(docx_file)

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level)
    if level == 0:  # 一级标题
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in heading.runs:
            run.font.size = Pt(22)
            run.font.bold = True
    elif level == 1:  # 二级标题
        for run in heading.runs:
            run.font.size = Pt(16)
            run.font.bold = True
    else:  # 三级标题
        for run in heading.runs:
            run.font.size = Pt(14)
            run.font.bold = True

def add_list_item(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text).font.name = '宋体'

def add_image_placeholder(doc, line):
    # 提取图片描述
    alt_text = re.search(r'!\[(.*?)\]', line).group(1)
    p = doc.add_paragraph()
    p.add_run(f"[图片: {alt_text}]").font.name = '宋体'
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']

if __name__ == '__main__':
    convert_md_to_docx('用户手册.md', '智能航运物流网络决策支持系统用户手册.docx')